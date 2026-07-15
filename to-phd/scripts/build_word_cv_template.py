import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


NAVY = "123F7A"
NAVY_DARK = "0E315F"
PALE_BLUE = "DCE5F0"
MID_BLUE = "8FA7C2"
TEXT = "222222"
MUTED = "5D6773"
WHITE = "FFFFFF"
FONT_CN = "微软雅黑"
FONT_LATIN = "Microsoft YaHei"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=30, start=50, bottom=30, end=50):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, data in edges.items():
        tag = "w:" + edge
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        for key, value in data.items():
            node.set(qn("w:" + key), str(value))


def remove_table_borders(table):
    nil = {"val": "nil"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=nil, left=nil, bottom=nil, right=nil, insideH=nil, insideV=nil)


def set_run_font(run, size=8.4, bold=False, color=TEXT, italic=False):
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, before=0, after=0, line=1.10, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_together = keep


def add_runs(paragraph, pieces, size=8.4):
    for piece in pieces:
        if isinstance(piece, str):
            piece = {"text": piece}
        run = paragraph.add_run(piece.get("text", ""))
        set_run_font(
            run,
            size=piece.get("size", size),
            bold=piece.get("bold", False),
            color=piece.get("color", TEXT),
            italic=piece.get("italic", False),
        )
    return paragraph


def section_header(doc, title):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Cm(2.9)
    table.columns[1].width = Cm(15.2)
    left, right = table.rows[0].cells
    left.width = Cm(2.9)
    right.width = Cm(15.2)
    set_cell_shading(left, NAVY)
    set_cell_shading(right, PALE_BLUE)
    set_cell_margins(left, top=46, start=70, bottom=46, end=70)
    set_cell_margins(right, top=46, start=55, bottom=46, end=55)
    for cell in (left, right):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border(cell, top={"val": "nil"}, left={"val": "nil"}, bottom={"val": "nil"}, right={"val": "nil"})
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(p, line=1)
    add_runs(p, [{"text": title, "bold": True, "color": WHITE, "size": 9.4}])
    p = right.paragraphs[0]
    format_paragraph(p, line=1)
    add_runs(p, [{"text": "◆", "color": MID_BLUE, "size": 7}])
    spacer = doc.add_paragraph()
    format_paragraph(spacer, line=0.90)


def add_inline(doc, label, value, size=8.1):
    p = doc.add_paragraph()
    format_paragraph(p, after=2.0, line=1.10, keep=True)
    add_runs(
        p,
        [
            {"text": label + "：", "bold": True, "color": NAVY_DARK, "size": size},
            {"text": value, "size": size},
        ],
    )


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.28)
    p.paragraph_format.first_line_indent = Cm(-0.28)
    format_paragraph(p, after=1.5, line=1.10, keep=True)
    add_runs(p, [{"text": "• " + text, "size": 7.9}])


def build(output_path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(0.9)
    section.bottom_margin = Cm(0.6)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    section.footer_distance = Cm(0.28)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(8.4)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.10

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(footer, [{"text": "姓名｜申博简历", "color": "8A929C", "size": 7}])

    title = doc.add_table(rows=1, cols=2)
    title.autofit = False
    title.columns[0].width = Cm(15.2)
    title.columns[1].width = Cm(2.9)
    remove_table_borders(title)
    left, right = title.rows[0].cells
    set_cell_margins(left, top=0, start=0, bottom=0, end=0)
    set_cell_margins(right, top=0, start=0, bottom=0, end=0)
    p = left.paragraphs[0]
    format_paragraph(p, after=1, line=1)
    add_runs(
        p,
        [
            {"text": "申博简历", "bold": True, "size": 20},
            {"text": "  PhD APPLICATION CV", "bold": True, "size": 10, "color": "30343A"},
        ],
    )
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_runs(p, [{"text": "◆◆◆", "bold": True, "size": 7.5, "color": NAVY}])

    bar = doc.add_table(rows=1, cols=2)
    bar.autofit = False
    bar.columns[0].width = Cm(13.2)
    bar.columns[1].width = Cm(4.9)
    for index, color in enumerate((NAVY, MID_BLUE)):
        cell = bar.rows[0].cells[index]
        set_cell_shading(cell, color)
        set_cell_margins(cell, top=22, start=0, bottom=22, end=0)
        set_cell_border(cell, top={"val": "nil"}, left={"val": "nil"}, bottom={"val": "nil"}, right={"val": "nil"})
        cell.paragraphs[0].clear()
    spacer = doc.add_paragraph()
    format_paragraph(spacer, line=0.90)

    section_header(doc, "个人信息")
    outer = doc.add_table(rows=1, cols=2)
    outer.autofit = False
    outer.alignment = WD_TABLE_ALIGNMENT.LEFT
    outer.columns[0].width = Cm(15.0)
    outer.columns[1].width = Cm(3.1)
    remove_table_borders(outer)
    info_cell, photo = outer.rows[0].cells
    set_cell_margins(info_cell, top=0, start=0, bottom=0, end=100)
    lead = info_cell.paragraphs[0]
    format_paragraph(lead, line=0.1)
    table = info_cell.add_table(rows=3, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(7.45)
    table.columns[1].width = Cm(7.45)
    remove_table_borders(table)
    fields = [
        ("姓　　名：", "[姓名]"),
        ("出生年月：", "[出生年月]"),
        ("毕业院校：", "[毕业院校]"),
        ("专　　业：", "[专业]"),
        ("联系电话：", "[联系电话]"),
        ("电子邮箱：", "[电子邮箱]"),
    ]
    for index, (label, value) in enumerate(fields):
        cell = table.cell(index // 2, index % 2)
        set_cell_margins(cell, top=16, start=0, bottom=16, end=20)
        p = cell.paragraphs[0]
        format_paragraph(p, line=1.06)
        add_runs(
            p,
            [
                {"text": label, "bold": True, "color": NAVY_DARK, "size": 8.3},
                {"text": value, "size": 8.3},
            ],
        )
    for row in table.rows:
        row.height = Cm(1.05)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    outer.rows[0].height = Cm(3.3)
    outer.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    photo.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_border(
        photo,
        top={"val": "single", "sz": "6", "color": "B6C0CC"},
        left={"val": "single", "sz": "6", "color": "B6C0CC"},
        bottom={"val": "single", "sz": "6", "color": "B6C0CC"},
        right={"val": "single", "sz": "6", "color": "B6C0CC"},
    )
    p = photo.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(p, [{"text": "证件照粘贴处", "color": "A0A8B2", "size": 7.3}])

    section_header(doc, "教育背景")
    education = doc.add_table(rows=2, cols=3)
    education.autofit = False
    education.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = (4.3, 7.4, 6.4)
    rows = [
        ("[起止时间]", "[硕士院校]", "[专业｜硕士]"),
        ("[起止时间]", "[本科院校]", "[专业｜本科]"),
    ]
    remove_table_borders(education)
    for row_index, row_data in enumerate(rows):
        for col_index, text in enumerate(row_data):
            cell = education.cell(row_index, col_index)
            cell.width = Cm(widths[col_index])
            set_cell_margins(cell, top=28, start=0, bottom=28, end=20)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if col_index == 2 else WD_ALIGN_PARAGRAPH.LEFT
            add_runs(p, [{"text": text, "bold": col_index in (1, 2), "size": 8.2}])

    section_header(doc, "研究经历")
    add_inline(doc, "硕士学位课题", "[课题名称]", size=8.2)
    add_bullet(doc, "[研究问题、数据或实验对象]。")
    add_bullet(doc, "[本人负责的实验、方法和分析工作]。")
    add_bullet(doc, "[当前真实进展与下一阶段计划；不要把计划写成已完成]。")

    section_header(doc, "论文成果")
    add_inline(doc, "1", "[作者]. [论文题目]. [期刊/会议], [年份]. [作者位置｜准确状态]", size=7.8)
    add_inline(doc, "2", "[作者]. [论文题目]. [目标期刊]. [作者位置｜准确状态]", size=7.8)
    add_inline(doc, "3", "[作者]. [论文题目]. [目标期刊]. [作者位置｜准确状态]", size=7.8)

    section_header(doc, "科研项目")
    add_inline(doc, "[项目来源]", "[项目名称]", size=7.9)
    add_bullet(doc, "[本人职责、方法和可公开结果]。")
    add_inline(doc, "[项目来源]", "[项目名称]", size=7.9)
    add_bullet(doc, "[本人职责、方法和可公开结果]。")
    add_inline(doc, "[项目来源]", "[项目名称]", size=7.9)
    add_bullet(doc, "[本人职责、方法和可公开结果]。")

    section_header(doc, "荣誉奖励")
    for text in (
        "[奖项名称]（[日期]）",
        "[奖学金或荣誉名称]（[日期]）",
        "[其他相关奖项]（[日期]）",
    ):
        p = doc.add_paragraph()
        format_paragraph(p, after=1.5, line=1.10, keep=True)
        add_runs(p, [{"text": text, "size": 8}])

    section_header(doc, "专业技能与英语")
    add_inline(doc, "研究与工具", "[软件、编程语言、实验与分析方法]", size=8)
    add_inline(doc, "英语能力", "[考试类型、分数、考试时间或待出分状态]", size=8)

    tail = doc.add_paragraph()
    tail.paragraph_format.space_before = Pt(0)
    tail.paragraph_format.space_after = Pt(0)
    tail.paragraph_format.line_spacing = Pt(1)
    set_run_font(tail.add_run(""), size=1)

    doc.core_properties.title = "申博简历Word模板"
    doc.core_properties.subject = "ToPhd中国大陆申博简历默认模板"
    doc.core_properties.author = "ToPhd"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def main():
    default_output = Path(__file__).resolve().parents[1] / "assets" / "申博简历-Word模板.docx"
    parser = argparse.ArgumentParser(description="生成ToPhd默认申博简历Word模板")
    parser.add_argument("--output", type=Path, default=default_output)
    args = parser.parse_args()
    print(str(build(args.output.resolve())))


if __name__ == "__main__":
    main()
